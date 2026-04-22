from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

out = '/mnt/data/deploy_site_professor/Relatorio_Deploy_Site_Estatico.docx'
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
style.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RELATÓRIO - DEPLOY DE SITE ESTÁTICO')
r.bold = True
r.font.size = Pt(15)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Trabalho acadêmico desenvolvido em HTML/CSS com proposta de publicação na Microsoft Azure.')

for title, body in [
    ('1. Objetivo',
     'Desenvolver uma página estática simples utilizando HTML e CSS e preparar sua publicação em nuvem, '
     'atendendo à proposta da atividade.'),
    ('2. Tecnologias utilizadas',
     '• HTML5 para a estrutura da página.\n'
     '• CSS3 para estilização e responsividade.\n'
     '• Microsoft Azure como plataforma sugerida para hospedagem.'),
    ('3. Estrutura do projeto',
     'O projeto foi organizado com dois arquivos principais:\n'
     '• index.html: contém toda a estrutura da página.\n'
     '• styles.css: contém a estilização visual do site.'),
    ('4. Descrição da página',
     'A página desenvolvida é um site institucional simples, contendo menu de navegação, seção de apresentação, '
     'seção sobre o projeto, área de serviços e contatos. O layout é responsivo e pode ser acessado em computador e celular.'),
    ('5. Processo de deploy no Azure Static Web Apps',
     '1. Criar uma conta no portal Microsoft Azure.\n'
     '2. Criar um recurso do tipo Static Web App.\n'
     '3. Enviar os arquivos index.html e styles.css para o serviço.\n'
     '4. Configurar a pasta principal do projeto como diretório de publicação.\n'
     '5. Após a finalização, copiar a URL gerada automaticamente pelo Azure.'),
    ('6. Alternativa de deploy no Azure Blob Storage',
     '1. Criar uma Storage Account no Azure.\n'
     '2. Ativar a opção Static website nas configurações da conta.\n'
     '3. Enviar os arquivos do site para o contêiner $web.\n'
     '4. Utilizar o endpoint disponibilizado pelo serviço para acessar o site publicado.'),
    ('7. Link do site',
     'Campo para preenchimento após a publicação no Azure:\n\n'
     'Link: ____________________________________________'),
    ('8. Conclusão',
     'O projeto atende ao objetivo solicitado, pois apresenta uma página estática funcional em HTML/CSS e está '
     'preparado para publicação em ambiente de nuvem por meio da plataforma Microsoft Azure.'),
]:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    for line in body.split('\n'):
        p = doc.add_paragraph(line)


doc.save(out)
print(out)
